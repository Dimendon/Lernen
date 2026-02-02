import pandas as pd
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
def ausfuehren(datum):
    doc = Document(r"D:\KI_Daten\Trainings_Anwesenheit\Anwesenheit1.docx")
    table = doc.tables[0]
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)
    rl = pd.DataFrame(data[1:], columns=data[0])
    rl.to_csv("Rangliste.csv", index=False)
    rl = rl.dropna(axis = 1)
    rl = rl.drop(["Su."],axis = 1)
    rl = rl.drop(["Jg/AGr"],axis = 1)
    rl = rl.drop([""],axis = 1)
    rl = rl.drop(["Nr."],axis = 1)
    rl["ÜN"] = pd.to_numeric(rl["ÜN"], errors="coerce")
    rangliste = rl.sort_values("ÜN", ascending = False)
    rangliste["Rang"] = range (1, len(rangliste) + 1)
    rangliste = rangliste.dropna(subset=["ÜN"])
    rangliste["Rang"] = rangliste["ÜN"].rank(
        method="dense",
        ascending=False
    ).astype(int)
    rangliste = rangliste[["Rang", "Vorname", "Name", "ÜN"]]
    print(rangliste)

    # Neues Word-Dokument
    DOCUMENT = "Rangliste.docx"
    document = docx.Document()

    u1 = document.add_heading("Erste Überschrift", 0)
    u2 = document.add_heading("Zweite Überschrift", 1)
    u3 = document.add_heading("Dritte Überschrift", 2)
    u1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(datum)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = document.add_table(rows=1, cols=len(rangliste.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(rangliste.columns):
        hdr_cells[i].text = str(col)

    for _, row in rangliste.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Speichern
    document.save("Rangliste.docx")
    return rangliste
#def datums():
    datums = eingabe.get()
    return datums

#GUI erstellen
root = tk.Tk()
root.title("Rangliste")
root.geometry("500x500")
#Text einfügen
text = tk.Label(root, text="Bitte das gewünschte Datum eingeben")
text.pack()
#Eingabefeld
xx = tk.Entry(root)
xx.pack()
x = 0

def eingeben():
    datum = xx.get()
    global x
    x = str(datum)

#Eingabe button
eingabe = tk.Button(
    root,
    text = "Eingabe",
    command = eingeben,

)
eingabe.pack()

#Escape Button
escape = tk.Button(
    root,
    text = "Verlassen",
    command = root.destroy
)
escape.pack()
root.mainloop()
#zweiter Schritt
ausfuehren(x)


